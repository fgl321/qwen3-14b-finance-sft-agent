from __future__ import annotations

import csv
import json
import re
from pathlib import Path
from typing import Any


SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".json",
    ".jsonl",
    ".csv",
    ".pdf",
    ".docx",
    ".png",
    ".jpg",
    ".jpeg",
}


def normalize_document_text(text: str) -> str:
    value = str(text).replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in value.split("\n")]
    result: list[str] = []
    blank = False
    for line in lines:
        clean = line.strip()
        if not clean:
            if result and not blank:
                result.append("")
            blank = True
            continue
        blank = False
        result.append(clean)
    return "\n".join(result).strip()


def _parse_json(path: Path) -> str:
    data = json.loads(path.read_text(encoding="utf-8-sig"))
    return json.dumps(data, ensure_ascii=False, indent=2)


def _parse_jsonl(path: Path) -> str:
    records: list[Any] = []
    for line in path.read_text(encoding="utf-8-sig").splitlines():
        if line.strip():
            records.append(json.loads(line))
    return "\n".join(
        json.dumps(item, ensure_ascii=False) for item in records
    )


def _parse_csv(path: Path) -> str:
    rows: list[str] = []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            rows.append(
                "；".join(f"{key}: {value}" for key, value in row.items())
            )
    return "\n".join(rows)


def _extract_pdf_pages(
    path: Path,
    *,
    ocr_enabled: bool = True,
) -> tuple[list[tuple[int, str]], int]:
    """
    逐页提取 PDF 文本。

    优先使用 PyMuPDF（fitz，C 实现，速度快且更抗坏页），
    回退到 pypdf；单页提取异常不会中断整份文档。
    返回 ([(page_number, text), ...], total_pages)。
    """
    total_pages = 0
    pages: list[tuple[int, str]] = []

    try:
        import pymupdf as fitz  # PyMuPDF >= 1.24
    except ImportError:
        try:
            import fitz  # 旧版本包名
        except ImportError:
            fitz = None

    if fitz is not None:
        try:
            document = fitz.open(str(path))
            try:
                total_pages = len(document)
                for index in range(total_pages):
                    try:
                        raw = document[index].get_text() or ""
                    except Exception:
                        raw = ""
                    text = normalize_document_text(raw)
                    if text:
                        pages.append((index + 1, text))
                        continue
                    if ocr_enabled:
                        try:
                            from app.rag.ocr import (
                                ocr_text_from_pixmap,
                            )

                            pixmap = document[index].get_pixmap(
                                matrix=fitz.Matrix(2, 2)
                            )
                            ocr_text = normalize_document_text(
                                ocr_text_from_pixmap(pixmap)
                            )
                            if ocr_text:
                                pages.append(
                                    (index + 1, ocr_text)
                                )
                        except Exception:
                            pass
            finally:
                document.close()
        except Exception:
            # PyMuPDF 打开失败时回退到 pypdf。
            pages = []
            total_pages = 0

    if pages or total_pages:
        return pages, total_pages

    from pypdf import PdfReader

    reader = PdfReader(str(path))
    total_pages = len(reader.pages)
    for index, page in enumerate(reader.pages, start=1):
        try:
            raw = page.extract_text() or ""
        except Exception:
            raw = ""
        text = normalize_document_text(raw)
        if text:
            pages.append((index, text))
        # pypdf 回退路径不做 OCR（无法渲染页面），
        # OCR 由 PyMuPDF 主路径覆盖。
    return pages, total_pages


def _looks_like_title_line(line: str) -> bool:
    candidate = line.strip()
    return bool(
        4 <= len(candidate) <= 60
        and re.search(r"[\u4e00-\u9fff]", candidate)
        and re.search(
            r"(?:条款|办法|规定|细则|指引|读本|手册|说明|报告|规则)",
            candidate,
        )
    )


def _merge_title_lines(lines: list[str]) -> str | None:
    """Merge a title line with continuation lines such as ``（第二版）``.

    Many Chinese PDFs split ``金融知识普及读本`` and ``（第二版）`` across two
    lines; keeping only the first line silently drops the edition marker.
    """

    for index, raw_line in enumerate(lines[:6]):
        line = raw_line.strip()
        if not _looks_like_title_line(line):
            continue
        merged = line
        for continuation in lines[index + 1 : index + 3]:
            piece = continuation.strip()
            if not piece:
                continue
            if re.match(r"^[（(]", piece) and re.search(r"[）)]$", piece):
                merged += piece
                continue
            break
        return merged
    return None


def _extract_pdf_title(
    path: Path,
    first_page_text: str,
) -> tuple[str | None, list[str]]:
    """Deterministic title extraction for document identity metadata.

    Prefers an explicit ``《...》`` book title on the first page, then PDF
    metadata title, then a short first line that looks like a document title.
    This only enriches Postgres metadata; it never decides document identity
    through fuzzy/vector matching.
    """

    file_name = path.name
    aliases: list[str] = []
    for value in (file_name, path.stem):
        if value and value not in aliases:
            aliases.append(value)

    metadata_title: str | None = None
    try:
        import pymupdf as fitz
    except ImportError:
        try:
            import fitz  # type: ignore[no-redef]
        except ImportError:
            fitz = None

    if fitz is not None:
        try:
            document = fitz.open(str(path))
            try:
                raw = (document.metadata or {}).get("title") or ""
                if raw.strip():
                    metadata_title = normalize_document_text(raw)
            finally:
                document.close()
        except Exception:
            metadata_title = None

    extracted: str | None = None
    if first_page_text:
        book_matches = re.findall(
            r"[《]([^》]{2,60})[》]",
            first_page_text,
        )
        if book_matches:
            extracted = book_matches[0].strip()
        else:
            candidate_lines = [
                line.strip()
                for line in first_page_text.splitlines()
                if line.strip()
            ][:6]
            extracted = _merge_title_lines(candidate_lines)

    if extracted is None and metadata_title and metadata_title != file_name:
        extracted = metadata_title

    for value in (extracted, metadata_title):
        if value and value not in aliases:
            aliases.append(value)
    return extracted, aliases


def _parse_pdf(path: Path, *, ocr_enabled: bool = True) -> str:
    pages, _ = _extract_pdf_pages(path, ocr_enabled=ocr_enabled)
    return "\n\n".join(text for _, text in pages)


def _parse_image(path: Path) -> str:
    from app.rag.ocr import ocr_text_from_bytes

    return ocr_text_from_bytes(path.read_bytes())

def _extract_docx_image_text(document: Any) -> list[str]:
    """提取 docx 内嵌图片并 OCR，返回识别文本（失败返回空列表）。"""
    from app.rag.ocr import ocr_text_from_bytes

    texts: list[str] = []
    try:
        rels = document.part.rels
    except Exception:
        return texts
    for rel in rels.values():
        try:
            target_ref = str(getattr(rel, "target_ref", "") or "").lower()
            if not target_ref.endswith(
                (".png", ".jpg", ".jpeg", ".gif", ".bmp")
            ):
                continue
            target_part = getattr(rel, "target_part", None)
            if target_part is None:
                continue
            data = bytes(target_part.blob)
            text = ocr_text_from_bytes(data)
            if text.strip():
                texts.append(f"[图片内容]\n{text}")
        except Exception:
            continue
    return texts


def _parse_docx(path: Path, *, ocr_enabled: bool = True) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - 依赖环境
        raise RuntimeError(
            "解析 DOCX 需要 python-docx：python -m pip install python-docx"
        ) from exc
    document = Document(str(path))
    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text for cell in row.cells))
    if ocr_enabled:
        blocks.extend(_extract_docx_image_text(document))
    return "\n".join(blocks)


def parse_document(path: str | Path) -> str:
    file_path = Path(path)
    if not file_path.exists() or not file_path.is_file():
        raise FileNotFoundError(str(file_path))
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"暂不支持文档类型 {suffix or '[无扩展名]'}。")

    if suffix in {".txt", ".md", ".markdown"}:
        text = file_path.read_text(encoding="utf-8-sig")
    elif suffix == ".json":
        text = _parse_json(file_path)
    elif suffix == ".jsonl":
        text = _parse_jsonl(file_path)
    elif suffix == ".csv":
        text = _parse_csv(file_path)
    elif suffix == ".pdf":
        text = _parse_pdf(file_path, ocr_enabled=True)
    elif suffix in {".png", ".jpg", ".jpeg"}:
        text = _parse_image(file_path)
    elif suffix == ".docx":
        text = _parse_docx(file_path, ocr_enabled=True)
    else:  # pragma: no cover
        raise ValueError(f"不支持的文档类型：{suffix}")

    normalized = normalize_document_text(text)
    if not normalized:
        raise ValueError("文档没有可提取的文本内容。")
    return normalized

# ---------------------------------------------------------------------------
# Legacy compatibility layer
# ---------------------------------------------------------------------------
# Stage 4.4 Lite uses the function-style parse_document() API above.  The
# original project also contains RagIngestionService, which imports and
# instantiates DocumentParser.  Keep both APIs so the new lifecycle service and
# the old knowledge upload route can coexist.

from dataclasses import asdict, dataclass, field
import tempfile
from collections.abc import Iterator


@dataclass(slots=True)
class ParsedPage:
    page_number: int
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def content(self) -> str:
        return self.text

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)

    def model_dump(self, *args: Any, **kwargs: Any) -> dict[str, Any]:
        del args, kwargs
        return self.to_dict()


class ParsedDocument(str):
    """String-compatible parsed result with legacy document attributes."""

    file_name: str
    source_type: str
    pages: list[ParsedPage]
    metadata: dict[str, Any]

    def __new__(
        cls,
        text: str,
        *,
        file_name: str,
        source_type: str,
        pages: list[ParsedPage] | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> "ParsedDocument":
        obj = str.__new__(cls, text)
        obj.file_name = file_name
        obj.source_type = source_type
        obj.pages = pages or [ParsedPage(page_number=1, text=text)]
        obj.metadata = dict(metadata or {})
        return obj

    @property
    def text(self) -> str:
        return str(self)

    @property
    def content(self) -> str:
        return str(self)

    @property
    def page_count(self) -> int:
        return len(self.pages)

    def iter_pages(self) -> Iterator[ParsedPage]:
        return iter(self.pages)

    def to_dict(self) -> dict[str, Any]:
        return {
            "file_name": self.file_name,
            "source_type": self.source_type,
            "text": str(self),
            "content": str(self),
            "page_count": self.page_count,
            "pages": [page.to_dict() for page in self.pages],
            "metadata": dict(self.metadata),
        }

    def model_dump(self, *args: Any, **kwargs: Any) -> dict[str, Any]:
        del args, kwargs
        return self.to_dict()


def _pages_from_path(path: Path, normalized_text: str) -> list[ParsedPage]:
    suffix = path.suffix.lower()
    if suffix != ".pdf":
        return [ParsedPage(page_number=1, text=normalized_text)]

    extracted, _ = _extract_pdf_pages(path)
    if not extracted:
        return [ParsedPage(page_number=1, text=normalized_text)]
    return [
        ParsedPage(page_number=page_number, text=text)
        for page_number, text in extracted
    ]


def _parse_document_with_pages(
    path: Path,
    *,
    ocr_enabled: bool = True,
) -> tuple[str, list[ParsedPage], dict[str, Any]]:
    """
    单次读取文档，返回 (全文, 页面列表, 元信息)。

    对 PDF 只读取一次（PyMuPDF/pypdf），避免旧实现重复解析两次；
    元信息中包含 total_pages 与 extracted_pages，便于上报
    因纯图片/无法提取而被跳过的页数。
    """
    suffix = path.suffix.lower()
    metadata: dict[str, Any] = {"source_path": str(path)}

    if suffix == ".pdf":
        extracted, total_pages = _extract_pdf_pages(
            path,
            ocr_enabled=ocr_enabled,
        )
        if not extracted:
            raise ValueError(
                "文档没有可提取的文本内容"
                "（纯图片/扫描件 PDF 未能识别出文字）。"
            )
        metadata["total_pages"] = total_pages
        metadata["extracted_pages"] = len(extracted)
        extracted_title, aliases = _extract_pdf_title(
            path,
            extracted[0][1],
        )
        if extracted_title:
            metadata["title"] = extracted_title
        metadata["aliases"] = aliases
        text = "\n\n".join(page_text for _, page_text in extracted)
        pages = [
            ParsedPage(page_number=page_number, text=page_text)
            for page_number, page_text in extracted
        ]
        return text, pages, metadata

    if suffix in {".png", ".jpg", ".jpeg"}:
        if not ocr_enabled:
            raise ValueError(
                "图片文档需要 OCR，但 OCR 功能未启用。"
            )
        text = _parse_image(path)
        if not text.strip():
            raise ValueError(
                "图片中没有识别到文字。"
            )
        metadata["source_type"] = "image"
        metadata["aliases"] = list(
            dict.fromkeys([path.name, path.stem])
        )
        return text, [ParsedPage(page_number=1, text=text)], metadata

    if suffix == ".docx":
        text = _parse_docx(path, ocr_enabled=ocr_enabled)
        metadata["aliases"] = list(
            dict.fromkeys([path.name, path.stem])
        )
        if not text.strip():
            raise ValueError("文档没有可提取的文本内容。")
        return text, [ParsedPage(page_number=1, text=text)], metadata

    text = parse_document(path)
    metadata["aliases"] = list(
        dict.fromkeys([path.name, path.stem])
    )
    return text, [ParsedPage(page_number=1, text=text)], metadata


class DocumentParser:
    """
    Backward-compatible parser facade.

    Supported calls include:
      parser.parse(Path("file.pdf"))
      parser.parse(file_path="file.pdf")
      parser.parse(file_name="file.txt", content=b"...")
      parser.parse_bytes(b"...", file_name="file.txt")

    The return value behaves like a normal string and also exposes legacy
    attributes such as .text, .content, .pages and .page_count.
    """

    supported_extensions = SUPPORTED_EXTENSIONS

    def __init__(self, settings: Any | None = None, **_: Any) -> None:
        self.settings = settings

    @staticmethod
    def _resolve_input(
        args: tuple[Any, ...], kwargs: dict[str, Any]
    ) -> tuple[Path | None, bytes | None, str | None]:
        path_value = (
            kwargs.pop("file_path", None)
            or kwargs.pop("path", None)
            or kwargs.pop("source_path", None)
        )
        file_name = (
            kwargs.pop("file_name", None)
            or kwargs.pop("filename", None)
            or kwargs.pop("name", None)
        )
        raw = (
            kwargs.pop("content", None)
            or kwargs.pop("file_bytes", None)
            or kwargs.pop("raw_bytes", None)
            or kwargs.pop("data", None)
        )

        if len(args) == 1:
            value = args[0]
            if isinstance(value, (bytes, bytearray, memoryview)):
                raw = bytes(value)
            else:
                path_value = value
        elif len(args) == 2:
            first, second = args
            if isinstance(first, (bytes, bytearray, memoryview)):
                raw = bytes(first)
                file_name = str(second)
            elif isinstance(second, (bytes, bytearray, memoryview)):
                file_name = str(first)
                raw = bytes(second)
            else:
                path_value = first
                file_name = str(second)
        elif len(args) > 2:
            raise TypeError("DocumentParser.parse 最多接受两个位置参数。")

        path = Path(path_value) if path_value is not None else None
        if raw is not None and not isinstance(raw, bytes):
            if isinstance(raw, str):
                raw = raw.encode("utf-8")
            else:
                raw = bytes(raw)
        return path, raw, str(file_name) if file_name else None

    def parse(self, *args: Any, **kwargs: Any) -> ParsedDocument:
        path, raw, file_name = self._resolve_input(args, dict(kwargs))
        cleanup_path: Path | None = None

        if raw is not None:
            final_name = file_name or "document.txt"
            suffix = Path(final_name).suffix.lower() or ".txt"
            if suffix not in SUPPORTED_EXTENSIONS:
                raise ValueError(f"暂不支持文档类型 {suffix}。")
            handle = tempfile.NamedTemporaryFile(
                prefix="finance_agent_parse_",
                suffix=suffix,
                delete=False,
            )
            try:
                handle.write(raw)
                handle.flush()
            finally:
                handle.close()
            path = Path(handle.name)
            cleanup_path = path
        elif path is None:
            raise TypeError("必须提供文档路径，或 file_name + content。")

        assert path is not None
        try:
            ocr_enabled = bool(
                getattr(
                    getattr(self, "settings", None),
                    "ocr_enabled",
                    True,
                )
            )
            text, pages, metadata = _parse_document_with_pages(
                path,
                ocr_enabled=ocr_enabled,
            )
            display_name = file_name or path.name
            source_type = path.suffix.lower().lstrip(".") or "text"
            return ParsedDocument(
                text,
                file_name=display_name,
                source_type=source_type,
                pages=pages,
                metadata=metadata,
            )
        finally:
            if cleanup_path is not None:
                cleanup_path.unlink(missing_ok=True)

    def parse_file(self, file_path: str | Path) -> ParsedDocument:
        return self.parse(file_path)

    def parse_path(self, file_path: str | Path) -> ParsedDocument:
        return self.parse(file_path)

    def parse_bytes(
        self,
        content: bytes,
        file_name: str = "document.txt",
    ) -> ParsedDocument:
        return self.parse(file_name=file_name, content=content)

    def parse_upload(
        self,
        file_name: str,
        content: bytes,
    ) -> ParsedDocument:
        return self.parse(file_name=file_name, content=content)

    def extract_text(self, *args: Any, **kwargs: Any) -> str:
        return str(self.parse(*args, **kwargs))

    def __call__(self, *args: Any, **kwargs: Any) -> ParsedDocument:
        return self.parse(*args, **kwargs)
